from flask import Flask, render_template, request, send_file
import os
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import speech_recognition as sr

load_dotenv()

app = Flask(__name__)
client = OpenAI(api_key=os.getenv("api_key"))

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def extract_text(filepath):
    if filepath.endswith(".pdf"):
        reader = PdfReader(filepath)
        return "".join(page.extract_text() for page in reader.pages)

    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)

    elif filepath.endswith(".wav"):
        recognizer = sr.Recognizer()
        with sr.AudioFile(filepath) as source:
            audio = recognizer.record(source)
        return recognizer.recognize_google(audio)

    return ""

def summarize(text, length):
    prompt = f"Summarize the text into {length} bullet key points."

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content

def save_pdf(summary):
    path = os.path.join(OUTPUT_FOLDER, "summary.pdf")
    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    elements = [Paragraph(summary, styles["Normal"])]
    doc.build(elements)
    return path

def save_docx(summary):
    path = os.path.join(OUTPUT_FOLDER, "summary.docx")
    doc = Document()
    doc.add_heading("AI Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(path)
    return path

@app.route("/", methods=["GET", "POST"])
def index():
    summary = ""

    if request.method == "POST":
        text = request.form.get("text")
        length = request.form.get("length")
        file = request.files.get("file")

        if file and file.filename != "":
            filepath = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(filepath)
            text = extract_text(filepath)

        if text:
            summary = summarize(text, length)
            save_pdf(summary)
            save_docx(summary)

    return render_template("index.html", summary=summary)

@app.route("/download/pdf")
def download_pdf():
    return send_file("outputs/summary.pdf", as_attachment=True)

@app.route("/download/docx")
def download_docx():
    return send_file("outputs/summary.docx", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
